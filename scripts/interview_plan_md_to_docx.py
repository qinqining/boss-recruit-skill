#!/usr/bin/env python3
"""将面试方案 Markdown 转为 Word（.docx）。用于满足「正式交付物为 .docx」。

用法（仓库根目录）：
  pip install python-docx
  python scripts/interview_plan_md_to_docx.py 面试方案生成任务/输出/面试方案_2026-05-13_张三.md

未传第二参数时，在同目录生成同名 .docx。
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("需要：pip install python-docx", file=sys.stderr)
    raise SystemExit(1) from None


def main() -> int:
    ap = argparse.ArgumentParser(description="Markdown 面试方案 → Word .docx（轻量解析）")
    ap.add_argument("md_path", type=Path, help="输入 .md 路径")
    ap.add_argument(
        "docx_path",
        type=Path,
        nargs="?",
        default=None,
        help="输出 .docx 路径（默认同名）",
    )
    args = ap.parse_args()
    md_path = args.md_path.resolve()
    if not md_path.is_file():
        print(f"文件不存在：{md_path}", file=sys.stderr)
        return 1
    docx_path = args.docx_path.resolve() if args.docx_path else md_path.with_suffix(".docx")

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    in_fence = False

    for raw in lines:
        line = raw.rstrip()
        if line.strip().startswith("```"):
            in_fence = not in_fence
            continue
        if in_fence:
            p = doc.add_paragraph(line, style="Normal")
            for r in p.runs:
                r.font.size = Pt(9)
            continue

        if not line.strip():
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif re.match(r"^[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", line).strip(), style="List Bullet")
        elif "|" in line and line.strip().startswith("|"):
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.left_indent = Pt(6)
            for r in p.runs:
                r.font.size = Pt(10)
        else:
            doc.add_paragraph(line.strip())

    doc.save(str(docx_path))
    print(f"已写入：{docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
